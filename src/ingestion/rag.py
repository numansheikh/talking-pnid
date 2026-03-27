"""
rag.py — Build RAG index from narrative DOCX files.

Chunks narrative documents, embeds with text-embedding-3-small,
saves chunks.json + embeddings.npy to data/rag/.

Usage:
  python rag.py          # index all docs (skips if index exists)
  python rag.py --force  # re-index everything

Output:
  src/talking-pnids-py/data/rag/chunks.json   — chunk text + metadata
  src/talking-pnids-py/data/rag/embeddings.npy — numpy float32 array (N x 1536)
"""

import json
import os
import sys
import time
from pathlib import Path

import numpy as np

sys.path.insert(0, str(Path(__file__).parent))
from config import REPO_ROOT, _load_key

# ── Paths ─────────────────────────────────────────────────────────────────────

NARRATIVES_DIR = REPO_ROOT / "data" / "datasets" / "rumaila-pp01" / "narratives"
RAG_OUT_DIR    = REPO_ROOT / "src" / "talking-pnids-py" / "data" / "rag"

# Each doc maps to the P&ID IDs it's relevant for.
# "global" = relevant to all P&IDs.
DOC_TAG_MAP = {
    "pid-006-narrative.docx":             ["pid-006"],
    "pid-007-narrative.docx":             ["pid-007"],
    "pid-008-narrative-a.docx":           ["pid-008"],
    "pid-008-narrative-b.docx":           ["pid-008"],
    "pid-008-detail-full.docx":           ["pid-008"],
    "detailed-doc-specification.docx":    ["global"],
    "scraper-receiver-details.docx":      ["pid-006", "pid-007"],
    "technical-spec-line-designation.docx": ["global"],
}

CHUNK_SIZE    = 1500   # characters (~375 tokens)
CHUNK_OVERLAP = 200    # characters
EMBED_MODEL   = "text-embedding-3-small"
EMBED_DIM     = 1536

# ── DOCX reading ──────────────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    """Extract plain text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        print("[rag] ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


# ── Chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks by character count."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        # Try to break at paragraph boundary
        if end < len(text):
            last_para = chunk.rfind("\n\n")
            last_sent = chunk.rfind(". ")
            break_at = max(last_para, last_sent)
            if break_at > size // 2:
                chunk = chunk[:break_at + 1]
                end = start + break_at + 1
        chunks.append(chunk.strip())
        start = end - overlap
    return [c for c in chunks if len(c) > 100]  # skip tiny trailing chunks


# ── Embedding ─────────────────────────────────────────────────────────────────

def _embed_batch(client, texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts, respecting API rate limits."""
    response = client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [item.embedding for item in response.data]


# ── Main indexing ─────────────────────────────────────────────────────────────

def build_index(force: bool = False) -> None:
    chunks_path = RAG_OUT_DIR / "chunks.json"
    embed_path  = RAG_OUT_DIR / "embeddings.npy"

    if chunks_path.exists() and embed_path.exists() and not force:
        existing = json.loads(chunks_path.read_text())
        print(f"[rag] Index already exists: {len(existing)} chunks. Use --force to rebuild.")
        return

    _load_key("apikey-openai-talking-pnid", "OPENAI_API_KEY")
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("[rag] ERROR: OPENAI_API_KEY not set")
        sys.exit(1)

    from openai import OpenAI
    client = OpenAI(api_key=api_key)

    RAG_OUT_DIR.mkdir(parents=True, exist_ok=True)

    all_chunks = []    # list of {text, source, pid_tags, chunk_index}
    all_texts  = []    # flat list of text strings for embedding

    print(f"[rag] Reading {len(DOC_TAG_MAP)} documents from {NARRATIVES_DIR}")

    for fname, pid_tags in DOC_TAG_MAP.items():
        fpath = NARRATIVES_DIR / fname
        if not fpath.exists():
            print(f"[rag]   SKIP (not found): {fname}")
            continue

        print(f"[rag]   Reading: {fname} → tags={pid_tags}")
        text = _read_docx(fpath)
        chunks = _chunk_text(text)
        print(f"[rag]     {len(text):,} chars → {len(chunks)} chunks")

        for i, chunk in enumerate(chunks):
            all_chunks.append({
                "text":        chunk,
                "source":      fname,
                "pid_tags":    pid_tags,
                "chunk_index": i,
            })
            all_texts.append(chunk)

    print(f"\n[rag] Total: {len(all_chunks)} chunks to embed")

    # Embed in batches of 100
    all_embeddings = []
    batch_size = 100
    t0 = time.time()

    for i in range(0, len(all_texts), batch_size):
        batch = all_texts[i:i + batch_size]
        print(f"[rag]   Embedding batch {i // batch_size + 1}/{(len(all_texts) + batch_size - 1) // batch_size} ({len(batch)} texts)...")
        embeddings = _embed_batch(client, batch)
        all_embeddings.extend(embeddings)

    elapsed = time.time() - t0
    embed_array = np.array(all_embeddings, dtype=np.float32)

    # Estimate cost: text-embedding-3-small = $0.02 / 1M tokens
    total_chars = sum(len(t) for t in all_texts)
    approx_tokens = total_chars // 4
    approx_cost = approx_tokens * 0.02 / 1_000_000

    print(f"\n[rag] Embedding complete: {len(all_chunks)} chunks, {embed_array.shape}")
    print(f"[rag] Approx tokens: {approx_tokens:,} | Cost: ~${approx_cost:.4f} | Time: {elapsed:.1f}s")

    # Save
    chunks_path.write_text(json.dumps(all_chunks, indent=2))
    np.save(str(embed_path), embed_array)

    print(f"[rag] Saved:")
    print(f"  → {chunks_path}  ({len(all_chunks)} chunks)")
    print(f"  → {embed_path}  (shape {embed_array.shape})")


if __name__ == "__main__":
    force = "--force" in sys.argv
    build_index(force=force)
