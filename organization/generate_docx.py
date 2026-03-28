"""
Generate three DOCX files:
  - brochure.docx         (4-section marketing document)
  - client-one-pager.docx (trial handout)
  - demo-script.docx      (internal rep guide)

Run: python3 generate_docx.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ───────────────────────────────────────────────────────────────────
TEAL      = RGBColor(0x0d, 0x94, 0x88)
DARK      = RGBColor(0x0f, 0x17, 0x2a)
NEAR_BLK  = RGBColor(0x1e, 0x29, 0x3b)
SLATE     = RGBColor(0x47, 0x55, 0x69)
LIGHT_GREY= RGBColor(0x94, 0xa3, 0xb8)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_para_border_bottom(para, color="0d9488", size=12):
    """Add a bottom border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_para_border_left(para, color="0d9488", size=24):
    """Add a left border (accent bar) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def h(doc, text, level=1, color=None, size=None, bold=None, align=None):
    """Add a styled heading."""
    defaults = {1: (24, True, NEAR_BLK), 2: (18, True, NEAR_BLK), 3: (13, True, TEAL)}
    sz, bd, col = defaults.get(level, (12, False, NEAR_BLK))
    if size:  sz  = size
    if bold is not None: bd = bold
    if color: col = color
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(sz)
    run.font.bold  = bd
    run.font.color.rgb = col
    return p


def body(doc, text, size=11, color=None, italic=False, align=None, space_after=6):
    """Add body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color or SLATE
    run.font.italic = italic
    return p


def bullet(doc, text, size=11, color=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color or SLATE
    return p


def spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)
    p.paragraph_format.space_before = Pt(0)


def teal_label(doc, text):
    """Uppercase teal section label."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size  = Pt(9)
    run.font.bold  = True
    run.font.color.rgb = TEAL
    run.font.all_caps  = True
    p.paragraph_format.space_after = Pt(2)
    return p


def page_break(doc):
    doc.add_page_break()


def set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


# ── BROCHURE.DOCX ─────────────────────────────────────────────────────────────
def make_brochure():
    doc = Document()
    set_margins(doc)

    # ── Cover section ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_border_bottom(p, color="0d9488", size=16)
    r = p.add_run("Talking P&IDs")
    r.font.size  = Pt(32)
    r.font.bold  = True
    r.font.color.rgb = TEAL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Your drawings. Finally answering back.")
    r2.font.size  = Pt(16)
    r2.font.color.rgb = NEAR_BLK
    r2.font.bold  = True

    spacer(doc, 4)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("AI-powered Q&A for industrial Piping & Instrumentation Diagrams")
    r3.font.size  = Pt(12)
    r3.font.color.rgb = SLATE
    r3.font.italic = True

    spacer(doc, 4)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("ZeerakLabs  ·  talkingpnid.zeeraklabs.online  ·  hello@zeeraklabs.com")
    r4.font.size  = Pt(10)
    r4.font.color.rgb = TEAL

    page_break(doc)

    # ── Section 1: The Problem ──
    teal_label(doc, "The Problem")
    h(doc, "Critical knowledge is buried in complex drawings", level=1)
    spacer(doc, 4)

    pb = doc.add_paragraph()
    set_para_border_left(pb, color="0d9488", size=18)
    pb.paragraph_format.left_indent = Inches(0.15)
    r = pb.add_run(
        "A modern process facility can have hundreds of P&ID drawings, each containing "
        "hundreds of instruments, valves, line segments, and interconnections. "
        "Engineers spend significant time — and facilities spend real money — on "
        "questions that should take seconds to answer."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = NEAR_BLK

    spacer(doc)

    pain_points = [
        ("Senior engineers as human search engines",
         "Experienced engineers are interrupted repeatedly to answer basic drawing questions."),
        ("Shutdown prep takes days",
         "Generating isolation valve lists, spec break summaries, and instrument checklists requires manual drawing review."),
        ("New engineers take months to get up to speed",
         "Understanding complex P&IDs requires sustained mentorship that is hard to compress."),
        ("Field decisions made without full information",
         "Engineers on the plant floor often can't access drawings easily, increasing the risk of errors."),
    ]

    for title, desc in pain_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(6)
        r1 = p.add_run("▸  " + title + "  —  ")
        r1.font.bold  = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = NEAR_BLK
        r2 = p.add_run(desc)
        r2.font.size  = Pt(11)
        r2.font.color.rgb = SLATE

    page_break(doc)

    # ── Section 2: The Solution ──
    teal_label(doc, "The Solution")
    h(doc, "Ask your P&IDs anything. Get engineering answers.", level=1)
    spacer(doc, 4)
    body(doc,
         "Talking P&IDs converts your existing drawings into a structured knowledge base, "
         "then lets every engineer on your team query it in plain English — from a browser, "
         "on any device, anywhere on the plant.",
         size=11, color=NEAR_BLK)
    spacer(doc)

    features = [
        ("Works on scanned PDFs",
         "No CAD files or special exports required. Your existing scanned P&ID PDFs are sufficient."),
        ("Understands topology, not just text",
         "The system models process flow and relationships — it knows what's upstream and downstream."),
        ("Reads your legend sheets",
         "Your specific abbreviations and valve symbols are learned from your own legend documents."),
        ("Cross-drawing awareness",
         "Questions spanning multiple interconnected P&IDs are answered by tracing connections automatically."),
        ("Conversational — no special syntax",
         "Engineers ask questions the way they'd ask a colleague. Follow-up questions work naturally."),
        ("Secure, isolated per client",
         "Each facility's drawings are processed and stored in an isolated environment."),
    ]

    for title, desc in features:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("✓  " + title + "  ")
        r1.font.bold  = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = TEAL
        r2 = p.add_run("— " + desc)
        r2.font.size  = Pt(11)
        r2.font.color.rgb = SLATE

    spacer(doc)
    teal_label(doc, "Sample questions engineers are asking today")
    spacer(doc, 2)

    sample_qs = [
        "What are all the isolation valves for vessel 362-V001?",
        "EZV-0002 closes unexpectedly — what are the upstream pressure effects and safeguards?",
        "List all locked-open and locked-closed valves on this drawing",
        "What instruments have high-high or low-low alarm functions?",
        "What are the design pressure and temperature for this vessel?",
        "Show me all spec breaks and their boundary equipment",
    ]
    for q in sample_qs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        set_para_border_left(p, color="0d9488", size=12)
        r = p.add_run(f'"{q}"')
        r.font.size   = Pt(11)
        r.font.italic = True
        r.font.color.rgb = NEAR_BLK

    page_break(doc)

    # ── Section 3: Implementation + Pilot ──
    teal_label(doc, "Implementation")
    h(doc, "Up and running on your drawings in days", level=1)
    spacer(doc, 4)
    body(doc,
         "There is no software to install and no changes to your existing drawing management systems. "
         "We process your P&IDs and give your team a private, secure web interface.",
         size=11, color=NEAR_BLK)
    spacer(doc)

    steps = [
        ("1. Send us your P&ID PDFs",
         "Scanned or digital PDFs in any format. We also take your legend and title block sheets "
         "to understand your specific notation."),
        ("2. We build the knowledge graph",
         "Our AI pipeline processes each drawing, extracting all components, connections, setpoints, "
         "valve positions, and relationships. Typically 3–5 working days for an initial set."),
        ("3. Validation with your team",
         "We run a structured review with your process engineers to verify accuracy on a representative "
         "sample of questions. Known gaps are documented and addressed."),
        ("4. Your team gets access",
         "A private, secure web application for your facility. Engineers log in from any browser "
         "on any device, including tablets in the field."),
    ]

    for title, desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(title)
        r1.font.bold  = True
        r1.font.size  = Pt(12)
        r1.font.color.rgb = NEAR_BLK
        body(doc, desc, size=11, color=SLATE, space_after=4)

    spacer(doc)

    # Pilot box via table
    teal_label(doc, "Structured Pilot Programme")
    spacer(doc, 2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "0d9488")
    cell.width = Inches(6.0)

    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cp.add_run(
        "We are offering a structured pilot to a small number of facilities. "
        "You provide a representative set of P&IDs; we build a working system on your data "
        "and work with your team to validate and improve it. "
        "Pricing is fixed for the pilot period with a clear path to full deployment.\n\n"
        "talkingpnid.zeeraklabs.online  ·  hello@zeeraklabs.com"
    )
    r.font.size  = Pt(11)
    r.font.color.rgb = WHITE
    r.font.bold  = False

    spacer(doc)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_border_bottom(p_footer, color="0d9488", size=8)
    rf = p_footer.add_run("ZeerakLabs  ·  zeeraklabs.com  ·  hello@zeeraklabs.com")
    rf.font.size  = Pt(9)
    rf.font.color.rgb = SLATE

    doc.save("brochure.docx")
    print("Saved → brochure.docx")


# ── CLIENT ONE-PAGER.DOCX ─────────────────────────────────────────────────────
def make_one_pager():
    doc = Document()
    set_margins(doc, top=0.8, bottom=0.8, left=1.0, right=1.0)

    # Header
    p = doc.add_paragraph()
    set_para_border_bottom(p, color="0d9488", size=16)
    r1 = p.add_run("Talking P&IDs  ")
    r1.font.size  = Pt(26)
    r1.font.bold  = True
    r1.font.color.rgb = TEAL
    r2 = p.add_run("— 48-Hour Trial Access")
    r2.font.size  = Pt(18)
    r2.font.bold  = True
    r2.font.color.rgb = NEAR_BLK

    p2 = doc.add_paragraph()
    r = p2.add_run("AI-powered Q&A for industrial P&ID diagrams  ·  by ZeerakLabs")
    r.font.size   = Pt(10)
    r.font.color.rgb = SLATE
    r.font.italic  = True
    spacer(doc, 6)

    # Hero paragraph
    ph = doc.add_paragraph()
    set_para_border_left(ph, color="0d9488", size=24)
    ph.paragraph_format.left_indent = Inches(0.15)
    ph.paragraph_format.space_after = Pt(10)
    r = ph.add_run(
        "Your P&IDs hold the critical knowledge your engineers need every day — isolation valve "
        "lineups, vessel conditions, instrument setpoints, spec breaks. Today that knowledge is "
        "locked inside complex drawings. Talking P&IDs makes it instantly accessible to every "
        "engineer on your team, in plain English."
    )
    r.font.size  = Pt(11)
    r.font.color.rgb = NEAR_BLK

    # Two-column table: What you can ask / How it works
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'

    c1 = tbl.rows[0].cells[0]
    c2 = tbl.rows[0].cells[1]

    def fill_col(cell, heading, items):
        p = cell.paragraphs[0]
        r = p.add_run(heading.upper())
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.color.rgb = TEAL
        r.font.all_caps = True
        for item in items:
            bp = cell.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(3)
            br = bp.add_run(item)
            br.font.size  = Pt(10)
            br.font.color.rgb = SLATE

    fill_col(c1, "What you can ask", [
        "Isolation valve lists for any vessel",
        "Design and operating conditions",
        "What happens if a valve closes or fails",
        "All locked-open / locked-closed valves",
        "Instrument alarm levels (HH/H/L/LL)",
        "Spec breaks with boundary equipment",
        "Note references and applicable lines",
        "Cross-P&ID system connections",
    ])

    fill_col(c2, "How it works", [
        "Works on scanned raster PDFs — no CAD needed",
        "Reads your legend sheets and standards",
        "Understands topology, not just text",
        "Traces flow across multiple P&IDs",
        "Conversational — follow-up questions work",
        "No special syntax required",
        "Web-based — tablet or laptop in the field",
    ])

    spacer(doc, 6)

    # Sample questions
    teal_label(doc, "Try these during your trial")
    spacer(doc, 2)
    qs = [
        "What are all the isolation valves for vessel 362-V001?",
        "What are the design pressure and temperature for 362-V001?",
        "EZV-0002 closes unexpectedly — what are the upstream effects?",
        "List all locked-open and locked-closed valves",
        "What instruments have HH or LL alarm functions?",
        "What are all the spec breaks on this P&ID?",
    ]
    for q in qs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.15)
        set_para_border_left(p, color="0d9488", size=10)
        r = p.add_run(f'"{q}"')
        r.font.size   = Pt(10)
        r.font.italic = True
        r.font.color.rgb = NEAR_BLK

    spacer(doc, 6)

    # Trial access box
    teal_label(doc, "Your trial access")
    spacer(doc, 2)
    tbl2 = doc.add_table(rows=2, cols=3)
    tbl2.style = 'Table Grid'
    headers = ["URL", "Username", "Password"]
    values  = ["talkingpnid.zeeraklabs.online", "pnid", "pakistan"]

    for i, (hdr, val) in enumerate(zip(headers, values)):
        h_cell = tbl2.rows[0].cells[i]
        set_cell_bg(h_cell, "0d9488")
        p = h_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.font.bold  = True
        r.font.size  = Pt(10)
        r.font.color.rgb = WHITE

        v_cell = tbl2.rows[1].cells[i]
        p2 = v_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(val)
        r2.font.size  = Pt(11)
        r2.font.bold  = True
        r2.font.color.rgb = NEAR_BLK

    spacer(doc, 6)

    # Footer
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_border_bottom(pf, color="0d9488", size=8)
    rf = pf.add_run(
        "Questions? hello@zeeraklabs.com  ·  ZeerakLabs AI Solutions  ·  zeeraklabs.com"
    )
    rf.font.size  = Pt(9)
    rf.font.color.rgb = SLATE

    doc.save("client-one-pager.docx")
    print("Saved → client-one-pager.docx")


# ── DEMO SCRIPT.DOCX ──────────────────────────────────────────────────────────
def make_demo_script():
    doc = Document()
    set_margins(doc)

    # Title
    p = doc.add_paragraph()
    set_para_border_bottom(p, color="0d9488", size=16)
    r = p.add_run("Talking P&IDs — Demo Script")
    r.font.size  = Pt(24)
    r.font.bold  = True
    r.font.color.rgb = TEAL
    spacer(doc, 2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Internal guide for ZeerakLabs representatives")
    r2.font.size   = Pt(11)
    r2.font.color.rgb = SLATE
    r2.font.italic  = True

    spacer(doc, 8)

    sections = [
        ("Before the meeting", [
            "Open the app at talkingpnid.zeeraklabs.online and log in",
            "Have the app on a screen the client can see — tablet or laptop works well",
            "Default diagram loaded: Fuel Gas KO Drum (PID-008) — use this as your primary demo",
            "Know your audience: process engineer, maintenance lead, or management? Adjust depth accordingly",
        ], None),

        ("Opening (2 min)", None,
         'Start with the problem, not the product:\n\n'
         '"How much time does your team spend hunting through P&ID drawings to answer questions '
         'that should take 30 seconds? A new operator asks a senior engineer — who\'s already busy — '
         'where the isolation valves are for a vessel. The engineer drops what they\'re doing, pulls '
         'up the drawing, traces the lines. Repeat that 10 times a day, across a 200-P&ID facility."\n\n'
         '"We built Talking P&IDs to give that time back."'),

        ("Scene 1 — The basics: finding equipment (2 min)", None,
         'Ask: "What are the design conditions for vessel 362-V001?"\n\n'
         'While loading: "This is a natural language question — no special syntax, no tag lookup, '
         'no searching. The system reads the diagram and pulls the data box from the vessel."\n\n'
         'Point out design pressure, temperature, operating conditions in the answer.'),

        ("Scene 2 — Isolation for maintenance (3 min)", None,
         'Ask: "What are all the isolation valves for vessel 362-V001?"\n\n'
         'While loading: "This is what takes a trained engineer 20–30 minutes to do from a cold '
         'start on an unfamiliar P&ID. You have to trace every line in and out of the vessel, '
         'identify each manual isolation, check the normal position."\n\n'
         'Follow-up: "Which of those are locked open?" — demonstrates conversational memory.'),

        ("Scene 3 — Troubleshooting scenario (3 min)", None,
         'Ask: "EZV-0002 has closed unexpectedly — what are the upstream pressure effects and what safeguards are in place?"\n\n'
         'While loading: "This requires understanding the topology — what\'s upstream, what\'s connected, '
         'what protection exists. The system traces the process flow and identifies the safeguards: PSVs, '
         'high-pressure trips, bypasses."\n\n'
         'Point out ESD logic and PSV references in the answer.'),

        ("Scene 4 — Cross-diagram question (2 min, if time permits)", None,
         'Switch to "All P&IDs" in the sidebar.\n\n'
         'Ask: "How does the DS-1 scraper receiver connect to the KO drum?"\n\n'
         '"We have three interconnected P&IDs loaded — two scraper receivers feeding a fuel gas '
         'knockout drum. The system understands the connections across all three drawings."'),

        ("Scene 5 — Hand it to them (2 min)", None,
         '"What would you actually want to know about one of your P&IDs? Ask it."\n\n'
         'Let them type a question. This is the most powerful moment in the demo.'),

        ("Closing (2 min)", None,
         '"What you\'ve seen works on scanned raster P&IDs — the same format most operating facilities have. '
         'No reformatting, no CAD files. We process the PDF directly."\n\n'
         '"The next step is running this on your drawings — we\'d need a set of P&IDs and a week of '
         'processing time to have a working system on your data."\n\n'
         '"We\'re offering a structured pilot to a small number of clients. You\'d get access to query '
         'your own P&IDs, we\'d work closely with your team, and you\'d see exactly where the system '
         'performs and where the gaps are."'),
    ]

    for title, bullets_list, prose in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        set_para_border_left(p, color="0d9488", size=20)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(title)
        r.font.size  = Pt(13)
        r.font.bold  = True
        r.font.color.rgb = NEAR_BLK

        if bullets_list:
            for b in bullets_list:
                bullet(doc, b)
        if prose:
            for para in prose.split('\n\n'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.3)
                r = p.add_run(para.strip())
                r.font.size  = Pt(11)
                r.font.color.rgb = SLATE

    spacer(doc, 10)

    # Q&A section
    p = doc.add_paragraph()
    set_para_border_bottom(p, color="0d9488", size=12)
    r = p.add_run("Common Questions and How to Handle Them")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = NEAR_BLK
    spacer(doc, 4)

    qas = [
        ('"How accurate is it?"',
         'On our benchmark set of 10 engineering questions against a real P&ID, we score 79/100. '
         'We know exactly where the gaps are and are actively closing them. For the 80% of '
         'day-to-day questions, the accuracy is high.'),
        ('"What data does it need?"',
         'Just the P&ID PDFs — scanned or digital. The drawings never leave a secure environment.'),
        ('"How long does onboarding take?"',
         'For 10–20 P&IDs, roughly a week of processing time.'),
        ('"Can it handle our legend and symbol standards?"',
         'Yes — it reads your legend sheets as part of the process and learns your specific conventions.'),
        ('"What about integration with our existing systems?"',
         'The current version is a standalone web app. API integration is on the roadmap.'),
    ]

    for q, a in qas:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(q + "  ")
        r1.font.bold  = True
        r1.font.size  = Pt(11)
        r1.font.color.rgb = NEAR_BLK
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Inches(0.25)
        r2 = p2.add_run(a)
        r2.font.size  = Pt(11)
        r2.font.color.rgb = SLATE

    doc.save("demo-script.docx")
    print("Saved → demo-script.docx")


# ── MAIN ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    make_brochure()
    make_one_pager()
    make_demo_script()
    print("\nAll done. Open in Word or upload to Google Docs.")
